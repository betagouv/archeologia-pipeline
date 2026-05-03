
#!/usr/bin/env python3
"""
Générateur de la documentation utilisateur Archéolog'IA
Produit un fichier .docx avec mise en forme professionnelle complète.

Prérequis : pip install python-docx
Usage : python generate_doc.py
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
import re

# ─── PALETTE DSFR (Bleu France / Rouge Marianne) ─────────────

# Primaire
BLEU_FONCE = RGBColor(0x00, 0x00, 0x91)        # Bleu France — titres H1, bandeaux
BLEU_ACCENT = RGBColor(0x6A, 0x6A, 0xF4)       # Bleu France clair — filets, accents
BLEU_SECONDAIRE = RGBColor(0x1F, 0x2E, 0x5E)   # Bleu nuit — H2/H3, bloc code

# Statuts
ROUGE = RGBColor(0xE1, 0x00, 0x0F)             # Rouge Marianne — alertes fortes
ORANGE = RGBColor(0xB3, 0x40, 0x00)            # Brun DSFR — avertissements
VERT = RGBColor(0x18, 0x75, 0x3C)              # Vert DSFR — puces & succès
VERT_FONCE = RGBColor(0x18, 0x75, 0x3C)
VERT_CLAIR_FOND = RGBColor(0xE3, 0xF5, 0xE1)
BLEU_NOTE_TEXTE = RGBColor(0x00, 0x63, 0xCB)   # Info DSFR — bloc Note
BLEU_NOTE_FOND = RGBColor(0xE8, 0xF1, 0xFA)
VIOLET_TEXTE = RGBColor(0x6E, 0x2B, 0x82)      # Accent Exemple
VIOLET_FOND = RGBColor(0xF4, 0xE8, 0xF8)

# Gris DSFR
GRIS_TEXTE = RGBColor(0x3A, 0x3A, 0x3A)
GRIS_SEC = RGBColor(0x66, 0x66, 0x66)
GRIS_FOND = RGBColor(0xF6, 0xF6, 0xF6)
GRIS_BORDURE = RGBColor(0xCE, 0xCE, 0xCE)
BLANC = RGBColor(0xFF, 0xFF, 0xFF)

# Fonds des callouts
BLEU_CLAIR_FOND = RGBColor(0xE3, 0xE3, 0xFD)   # Fond Conseil (bleu France très clair)
JAUNE_CLAIR_FOND = RGBColor(0xFD, 0xF4, 0xE5)  # Fond Avertissement
JAUNE_TEXTE = RGBColor(0x71, 0x50, 0x0D)       # Texte Avertissement

# Code
FOND_CODE = RGBColor(0x1F, 0x2E, 0x5E)
TEXTE_CODE = RGBColor(0xEC, 0xF0, 0xF1)
CODE_INLINE_ROUGE = RGBColor(0xE1, 0x00, 0x0F)

# ─── TYPOGRAPHIE ───────────────────────────────────────
FONT_MAIN = "Segoe UI"           # Police de base (titres + corps)
FONT_MONO = "Consolas"            # Police monospace (code)

# ─── FONCTIONS UTILITAIRES ──────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Applique une couleur de fond à une cellule de tableau."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="DEE2E6", size="4"):
    """Applique des bordures à une cellule."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def add_horizontal_line(doc, color="000091", width=60):
    """Ajoute un filet horizontal décoratif."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def add_heading1(doc, text, bookmark=None):
    """Titre de niveau 1 (style Word natif « Heading 1 »).

    Ornement DSFR : filet vertical Bleu France à gauche du titre.
    Utiliser le style natif permet à Word de générer la TDM automatique
    et d'afficher le titre dans le volet Navigation.

    `bookmark` : si fourni, pose un signet Word nommé ainsi autour du
    titre, pour servir de cible aux hyperliens internes.
    """
    p = doc.add_paragraph(style="Heading 1")
    # Ornement : filet gauche Bleu France (large)
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="36" w:space="10" w:color="000091"/>'
        f'</w:pBdr>'
    ))
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = BLEU_FONCE
    run.font.name = FONT_MAIN
    if bookmark:
        _add_bookmark_around_paragraph(p, bookmark)
    add_horizontal_line(doc, color="6A6AF4")
    return p


def add_heading2(doc, text):
    """Titre de niveau 2 (style Word natif « Heading 2 »)."""
    p = doc.add_paragraph(style="Heading 2")
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = BLEU_SECONDAIRE
    run.font.name = FONT_MAIN
    return p


def add_heading3(doc, text):
    """Titre de niveau 3 (style Word natif « Heading 3 »)."""
    p = doc.add_paragraph(style="Heading 3")
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = BLEU_SECONDAIRE
    run.font.name = FONT_MAIN
    return p


def add_body(doc, text):
    """Ajoute un paragraphe de corps de texte."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = GRIS_TEXTE
    run.font.name = FONT_MAIN
    return p


def add_body_rich(doc, segments):
    """
    Ajoute un paragraphe avec segments riches.
    segments = [(texte, bold, italic, color, font_name, font_size), ...]
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    for seg in segments:
        text = seg[0]
        bold = seg[1] if len(seg) > 1 else False
        italic = seg[2] if len(seg) > 2 else False
        color = seg[3] if len(seg) > 3 else GRIS_TEXTE
        fname = seg[4] if len(seg) > 4 else FONT_MAIN
        fsize = seg[5] if len(seg) > 5 else Pt(11)
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = color
        run.font.name = fname
        run.font.size = fsize
    return p


def add_bullet(doc, text, level=0):
    """Ajoute un élément de liste à puces."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    bullet = "•" if level == 0 else "◦"
    run = p.add_run(f"{bullet}  {text}")
    run.font.size = Pt(11)
    run.font.color.rgb = GRIS_TEXTE
    run.font.name = FONT_MAIN
    return p


def add_conseil(doc, text):
    """Ajoute un bloc Conseil stylé."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "E3E3FD")

    # Bordure gauche bleue
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="E3E3FD"/>'
        f'  <w:left w:val="single" w:sz="18" w:space="0" w:color="000091"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="E3E3FD"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="E3E3FD"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

    p = cell.paragraphs[0]
    run = p.add_run("💡 Conseil — ")
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLEU_ACCENT
    run.font.name = FONT_MAIN
    run2 = p.add_run(text)
    run2.font.size = Pt(10.5)
    run2.font.color.rgb = BLEU_FONCE
    run2.font.name = FONT_MAIN
    doc.add_paragraph()  # Espace après


def add_avertissement(doc, text):
    """Ajoute un bloc Avertissement stylé."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FDF4E5")

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="FDF4E5"/>'
        f'  <w:left w:val="single" w:sz="18" w:space="0" w:color="B34000"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="FDF4E5"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="FDF4E5"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

    p = cell.paragraphs[0]
    run = p.add_run("⚠️ Attention — ")
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = ORANGE
    run.font.name = FONT_MAIN
    run2 = p.add_run(text)
    run2.font.size = Pt(10.5)
    run2.font.color.rgb = JAUNE_TEXTE
    run2.font.name = FONT_MAIN
    doc.add_paragraph()


def add_styled_table(doc, headers, rows, col_widths=None):
    """Ajoute un tableau stylé avec en-tête bleu foncé et lignes alternées."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # En-tête
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, "000091")
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = BLANC
        run.font.name = FONT_MAIN

    # Lignes
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx + 1, c_idx)
            if r_idx % 2 == 0:
                set_cell_shading(cell, "F8F9FA")
            else:
                set_cell_shading(cell, "FFFFFF")
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.color.rgb = GRIS_TEXTE
            run.font.name = FONT_MAIN

    # Largeurs de colonnes
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # Espace après
    return table


def add_code_block(doc, code_text):
    """Ajoute un bloc de code stylé fond sombre."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "1F2E5E")

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="1F2E5E"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="1F2E5E"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="1F2E5E"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="1F2E5E"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

    for i, line in enumerate(code_text.strip().split("\n")):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        run = p.add_run(line)
        run.font.size = Pt(9.5)
        run.font.color.rgb = TEXTE_CODE
        run.font.name = "Consolas"
    doc.add_paragraph()


def add_screenshot_placeholder(doc, caption):
    """Ajoute un espace réservé pour une capture d'écran."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8F9FA")
    set_cell_borders(cell, color="DEE2E6")

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n[ CAPTURE D'ÉCRAN ]\n")
    run.font.size = Pt(11)
    run.font.color.rgb = GRIS_SEC
    run.font.name = FONT_MAIN
    run.italic = True

    # Légende
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(12)
    run2 = p2.add_run(caption)
    run2.font.size = Pt(9.5)
    run2.font.color.rgb = GRIS_SEC
    run2.font.name = FONT_MAIN
    run2.italic = True


def add_page_break(doc):
    doc.add_page_break()


# ─── CHAMPS DYNAMIQUES (pagination, table des matières) ─────

def _insert_field(paragraph, instr_text, dirty=False):
    """Insère un champ Word (PAGE, NUMPAGES, TOC, SEQ, …) dans `paragraph`.

    Structure OOXML correcte : chaque fldChar et instrText est dans son
    propre <w:r>. C'est indispensable pour que Word puisse remplir la zone
    'separate' → 'end' avec son contenu dynamique (hyperliens pour TOC,
    valeurs pour STYLEREF, numéros pour SEQ, etc.).

    Retourne le run portant l'instrText, afin que l'appelant puisse
    styliser la police (Word applique ces propriétés à la valeur rendue).
    """
    # Run 1 : begin
    r_begin = paragraph.add_run()
    el_begin = OxmlElement("w:fldChar")
    el_begin.set(qn("w:fldCharType"), "begin")
    if dirty:
        el_begin.set(qn("w:dirty"), "true")
    r_begin._r.append(el_begin)

    # Run 2 : instrText (instruction du champ)
    r_instr = paragraph.add_run()
    el_instr = OxmlElement("w:instrText")
    el_instr.set(qn("xml:space"), "preserve")
    el_instr.text = instr_text
    r_instr._r.append(el_instr)

    # Run 3 : separate (délimite début de la valeur calculée)
    r_sep = paragraph.add_run()
    el_sep = OxmlElement("w:fldChar")
    el_sep.set(qn("w:fldCharType"), "separate")
    r_sep._r.append(el_sep)

    # Run 4 : valeur cachée par défaut (placeholder).
    # Certaines versions de Word affichent le code du champ à la place de
    # sa valeur lorsque la zone 'separate' → 'end' est totalement vide.
    # Ce run vide (ou ici '0') sert de slot qui sera remplacé par la vraie
    # valeur lors de la mise à jour des champs.
    paragraph.add_run("0")

    # Run 5 : end
    r_end = paragraph.add_run()
    el_end = OxmlElement("w:fldChar")
    el_end.set(qn("w:fldCharType"), "end")
    r_end._r.append(el_end)

    return r_instr


# Identifiants de signets (incrémentés à chaque bookmark créé)
_BOOKMARK_ID_COUNTER = [1000]


def _add_bookmark_around_paragraph(paragraph, name):
    """Entoure le contenu du paragraphe d'un signet Word nommé `name`.
    Ce signet sert de cible aux hyperliens internes.
    """
    bid = str(_BOOKMARK_ID_COUNTER[0])
    _BOOKMARK_ID_COUNTER[0] += 1
    p = paragraph._p
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bid)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bid)
    pPr = p.find(qn("w:pPr"))
    if pPr is not None:
        pPr.addnext(start)
    else:
        p.insert(0, start)
    p.append(end)


def _add_internal_hyperlink(paragraph, text, bookmark_name, *,
                            color_hex=None, bold=False, italic=False,
                            size_pt=None, font_name=None, underline=True):
    """Ajoute un hyperlien interne (vers un signet) dans `paragraph`."""
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    hyperlink.set(qn("w:history"), "1")
    run_el = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if font_name:
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)
        rPr.append(rFonts)
    if bold:
        rPr.append(OxmlElement("w:b"))
    if italic:
        rPr.append(OxmlElement("w:i"))
    if size_pt is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size_pt * 2)))
        rPr.append(sz)
    if color_hex is not None:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color_hex)
        rPr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    run_el.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run_el.append(t)
    hyperlink.append(run_el)
    paragraph._p.append(hyperlink)
    return hyperlink


def _small_italic_run(paragraph, text, color=None):
    run = paragraph.add_run(text)
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = color or GRIS_SEC
    run.font.name = FONT_MAIN
    return run


def configure_heading_styles(doc):
    """Configure les styles natifs Word (titres + Normal) pour garantir
    un rendu homogène et une TDM automatique fonctionnelle.
    """
    def _cfg(style_name, size_pt, color, sb_pt, sa_pt):
        s = doc.styles[style_name]
        s.font.name = FONT_MAIN
        s.font.bold = True
        s.font.size = Pt(size_pt)
        s.font.color.rgb = color
        pf = s.paragraph_format
        pf.space_before = Pt(sb_pt)
        pf.space_after = Pt(sa_pt)
        pf.keep_with_next = True

    _cfg("Heading 1", 22, BLEU_FONCE, 30, 6)
    _cfg("Heading 2", 16, BLEU_SECONDAIRE, 20, 8)
    _cfg("Heading 3", 13, BLEU_SECONDAIRE, 14, 6)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_MAIN
    normal.font.size = Pt(11)
    normal.font.color.rgb = GRIS_TEXTE
    normal.paragraph_format.line_spacing = 1.35


def setup_page_header_footer(doc, title="Archéolog'IA — Documentation utilisateur",
                              footer_left="Archéolog'IA"):
    """Ajoute un en-tête et un pied de page dans chaque section.

    En-tête : `<titre doc>`    TAB    `<nom de la section courante via STYLEREF>`
    Pied de page : `<texte gauche>`    TAB    `Page X / Y`
    La page de garde n'hérite ni de l'en-tête ni du pied de page.
    """
    for section in doc.sections:
        section.different_first_page_header_footer = True

        # ── En-tête (titre doc à gauche, nom de section à droite) ──
        hp = section.header.paragraphs[0]
        hp.paragraph_format.space_after = Pt(4)
        section_width = section.page_width - section.left_margin - section.right_margin
        hp.paragraph_format.tab_stops.add_tab_stop(section_width, WD_TAB_ALIGNMENT.RIGHT)

        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_title = hp.add_run(title)
        r_title.bold = True
        r_title.italic = True
        r_title.font.size = Pt(9)
        r_title.font.color.rgb = GRIS_SEC
        r_title.font.name = FONT_MAIN

        pPr = hp._p.get_or_add_pPr()
        pPr.append(parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="DEE2E6"/>'
            f'</w:pBdr>'
        ))

        # ── Pied de page ──
        fp = section.footer.paragraphs[0]
        fp.paragraph_format.space_before = Pt(4)
        fp.paragraph_format.tab_stops.add_tab_stop(section_width, WD_TAB_ALIGNMENT.RIGHT)
        _small_italic_run(fp, footer_left)
        fp.add_run("\t")
        _small_italic_run(fp, "Page ")
        r_page = _insert_field(fp, " PAGE ")
        r_page.font.size = Pt(9)
        r_page.italic = True
        r_page.font.color.rgb = GRIS_SEC
        r_page.font.name = FONT_MAIN
        _small_italic_run(fp, " / ")
        r_total = _insert_field(fp, " NUMPAGES ")
        r_total.font.size = Pt(9)
        r_total.italic = True
        r_total.font.color.rgb = GRIS_SEC
        r_total.font.name = FONT_MAIN


# ─── PAGE DE COUVERTURE ─────────────────────────────

def _cover_band(doc, color_hex, height_pt=10):
    """Bande horizontale colorée pleine largeur (page de couverture)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    ))
    r = p.add_run()
    r.font.size = Pt(height_pt)
    return p


def add_cover_page(doc, version="Version 2.0", organisation="betagouv · 2026"):
    """Page de garde raffinée : bandeaux DSFR, titre, filet d'accent, sous-titre."""
    # Bandeau DSFR en haut (bleu France + rouge marianne)
    _cover_band(doc, "000091", height_pt=12)
    _cover_band(doc, "E1000F", height_pt=8)

    for _ in range(4):
        doc.add_paragraph()

    # Petite étiquette d'identité
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PLUGIN QGIS")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = BLEU_ACCENT
    r.font.name = FONT_MAIN

    # Titre principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run("Archéolog'IA")
    r.bold = True
    r.font.size = Pt(48)
    r.font.color.rgb = BLEU_FONCE
    r.font.name = FONT_MAIN

    # Filet central (bandeau court)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Cm(5.5)
    p.paragraph_format.right_indent = Cm(5.5)
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="24" w:space="4" w:color="000091"/>'
        f'</w:pBdr>'
    ))

    # Sous-titre
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Documentation utilisateur")
    r.italic = True
    r.font.size = Pt(22)
    r.font.color.rgb = BLEU_SECONDAIRE
    r.font.name = FONT_MAIN

    # Description
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run(
        "Détection archéologique par traitement LiDAR\n"
        "et vision artificielle"
    )
    r.font.size = Pt(13)
    r.font.color.rgb = GRIS_SEC
    r.font.name = FONT_MAIN

    # Bas de page : version + organisation
    for _ in range(7):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(version)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = BLEU_FONCE
    r.font.name = FONT_MAIN

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(organisation)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = GRIS_SEC
    r.font.name = FONT_MAIN

    for _ in range(2):
        doc.add_paragraph()
    # Bandeau DSFR en bas (miroir du haut)
    _cover_band(doc, "E1000F", height_pt=8)
    _cover_band(doc, "000091", height_pt=12)

    add_page_break(doc)


# ─── TABLE DES MATIÈRES AUTOMATIQUE ─────────────────────────

def add_auto_toc(doc):
    """Insère une TDM automatique (niveaux 1 à 3) sous forme de champ TOC
    cliquable. L'utilisateur met à jour les champs à l'ouverture dans Word
    (clic droit → Mettre à jour les champs, ou touche F9).
    """
    add_heading1(doc, "Table des matières")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    _small_italic_run(
        p,
        "Astuce : si les numéros de page ne s'affichent pas, faites un clic droit "
        "dans la TDM et choisissez « Mettre à jour les champs » (ou appuyez sur F9).",
    )

    p = doc.add_paragraph()
    _insert_field(p, ' TOC \\o "1-3" \\h \\z \\u ', dirty=True)


# ─── CALLOUTS SPÉCIALISÉS ───────────────────────────────────

def _styled_callout(doc, icon, fill_hex, border_hex, icon_color, text_color, text):
    """Mise en forme commune des blocs Conseil / Avertissement / Note / Succès :
    fond coloré, large bordure gauche, titre en gras suivi du corps.
    """
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill_hex)
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{fill_hex}"/>'
        f'  <w:left w:val="single" w:sz="18" w:space="0" w:color="{border_hex}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{fill_hex}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{fill_hex}"/>'
        f'</w:tcBorders>'
    ))
    p = cell.paragraphs[0]
    r1 = p.add_run(icon)
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = icon_color
    r1.font.name = FONT_MAIN
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = text_color
    r2.font.name = FONT_MAIN
    doc.add_paragraph()


def add_note(doc, text):
    """Bloc d'information neutre (bleu clair)."""
    _styled_callout(doc, icon="ℹ️ Note — ", fill_hex="E8F1FA", border_hex="0063CB",
                    icon_color=BLEU_NOTE_TEXTE, text_color=BLEU_NOTE_TEXTE, text=text)


def add_succes(doc, text):
    """Bloc « à retenir / bonne pratique » (vert)."""
    _styled_callout(doc, icon="✅ À retenir — ", fill_hex="E3F5E1", border_hex="18753C",
                    icon_color=VERT_FONCE, text_color=VERT_FONCE, text=text)


# ─── CALLOUT D'ÉTAPE NUMÉROTÉE (Quick Start) ────────────────

def _no_borders_xml():
    return parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="nil"/><w:left w:val="nil"/>'
        f'  <w:bottom w:val="nil"/><w:right w:val="nil"/>'
        f'</w:tcBorders>'
    )


def add_step_callout(doc, number, title):
    """Bloc d'étape numéroté : grosse bulle bleue à gauche, titre à droite.
    Les étapes suivent le style Word « Heading 2 » (via un paragraphe
    vide invisible) afin d'apparaître dans la TDM automatique.
    """
    # Paragraphe fantôme en Heading 2 pour que l'étape apparaisse dans la TDM
    hp = doc.add_paragraph(style="Heading 2")
    hp.paragraph_format.space_before = Pt(14)
    hp.paragraph_format.space_after = Pt(0)
    hrun = hp.add_run(f"Étape {number} — {title}")
    hrun.bold = True
    hrun.font.size = Pt(1)  # invisible
    hrun.font.color.rgb = BLANC

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    # Cellule numéro (pastille bleue, compacte)
    num_cell = table.cell(0, 0)
    num_cell.width = Cm(1.0)
    set_cell_shading(num_cell, "000091")
    num_cell._tc.get_or_add_tcPr().append(_no_borders_xml())
    np = num_cell.paragraphs[0]
    np.alignment = WD_ALIGN_PARAGRAPH.CENTER
    np.paragraph_format.space_before = Pt(2)
    np.paragraph_format.space_after = Pt(2)
    r = np.add_run(str(number))
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = BLANC
    r.font.name = FONT_MAIN

    # Cellule titre
    title_cell = table.cell(0, 1)
    title_cell.width = Cm(14.5)
    tc2 = title_cell._tc.get_or_add_tcPr()
    tc2.append(parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="nil"/><w:left w:val="nil"/>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="000091"/>'
        f'  <w:right w:val="nil"/>'
        f'</w:tcBorders>'
    ))
    tp = title_cell.paragraphs[0]
    tp.paragraph_format.space_before = Pt(8)
    tp.paragraph_format.left_indent = Cm(0.3)
    tr = tp.add_run(f"Étape {number} — {title}")
    tr.bold = True
    tr.font.size = Pt(15)
    tr.font.color.rgb = BLEU_FONCE
    tr.font.name = FONT_MAIN

    doc.add_paragraph()


# ─── BLOC EXEMPLE (violet) ──────────────────────────────

def add_example(doc, text):
    """Bloc « Exemple » (violet) pour illustrer un cas concret."""
    _styled_callout(doc, icon="📌 Exemple — ", fill_hex="F4E8F8", border_hex="6E2B82",
                    icon_color=VIOLET_TEXTE, text_color=VIOLET_TEXTE, text=text)


# ─── BLOC PROCÉDURE (mini-étapes inline) ───────────────

def add_procedure(doc, steps):
    """Procédure pas-à-pas compacte : une ligne par étape numérotée.
    `steps` est une liste de chaînes.
    """
    table = doc.add_table(rows=len(steps), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for i, step in enumerate(steps):
        # Numéro (pastille claire)
        num_cell = table.cell(i, 0)
        num_cell.width = Cm(0.9)
        set_cell_shading(num_cell, "E3E3FD")
        num_cell._tc.get_or_add_tcPr().append(_no_borders_xml())
        np = num_cell.paragraphs[0]
        np.alignment = WD_ALIGN_PARAGRAPH.CENTER
        np.paragraph_format.space_before = Pt(2)
        np.paragraph_format.space_after = Pt(2)
        r = np.add_run(str(i + 1))
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = BLEU_FONCE
        r.font.name = FONT_MAIN

        # Texte de l'étape
        txt_cell = table.cell(i, 1)
        txt_cell.width = Cm(15)
        set_cell_shading(txt_cell, "F6F6FE")
        txt_cell._tc.get_or_add_tcPr().append(_no_borders_xml())
        tp = txt_cell.paragraphs[0]
        tp.paragraph_format.space_before = Pt(3)
        tp.paragraph_format.space_after = Pt(3)
        tp.paragraph_format.left_indent = Cm(0.2)
        tr = tp.add_run(step)
        tr.font.size = Pt(10.5)
        tr.font.color.rgb = GRIS_TEXTE
        tr.font.name = FONT_MAIN

    doc.add_paragraph()


# ─── CARTE DÉPANNAGE (Q / R) ───────────────────────────

def add_depannage_card(doc, question, solutions):
    """Carte Dépannage : question en en-tête + liste de solutions en corps."""
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # En-tête : question
    q_cell = table.cell(0, 0)
    set_cell_shading(q_cell, "FEE9E9")
    tcPr = q_cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="E1000F"/>'
        f'  <w:left w:val="single" w:sz="6" w:space="0" w:color="E1000F"/>'
        f'  <w:bottom w:val="single" w:sz="2" w:space="0" w:color="F5C2C3"/>'
        f'  <w:right w:val="single" w:sz="6" w:space="0" w:color="E1000F"/>'
        f'</w:tcBorders>'
    ))
    qp = q_cell.paragraphs[0]
    qp.paragraph_format.space_before = Pt(4)
    qp.paragraph_format.space_after = Pt(4)
    r_icon = qp.add_run("❓  ")
    r_icon.font.size = Pt(11)
    r_q = qp.add_run(question)
    r_q.bold = True
    r_q.font.size = Pt(11)
    r_q.font.color.rgb = ROUGE
    r_q.font.name = FONT_MAIN

    # Corps : solutions
    s_cell = table.cell(1, 0)
    set_cell_shading(s_cell, "FFFFFF")
    tcPr2 = s_cell._tc.get_or_add_tcPr()
    tcPr2.append(parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="2" w:space="0" w:color="F5C2C3"/>'
        f'  <w:left w:val="single" w:sz="6" w:space="0" w:color="E1000F"/>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="E1000F"/>'
        f'  <w:right w:val="single" w:sz="6" w:space="0" w:color="E1000F"/>'
        f'</w:tcBorders>'
    ))
    first = True
    for sol in solutions:
        if first:
            sp = s_cell.paragraphs[0]
            first = False
        else:
            sp = s_cell.add_paragraph()
        sp.paragraph_format.space_before = Pt(2)
        sp.paragraph_format.space_after = Pt(2)
        sp.paragraph_format.left_indent = Cm(0.3)
        r_arrow = sp.add_run("➜  ")
        r_arrow.bold = True
        r_arrow.font.size = Pt(10.5)
        r_arrow.font.color.rgb = ROUGE
        r_arrow.font.name = FONT_MAIN
        r_txt = sp.add_run(sol)
        r_txt.font.size = Pt(10.5)
        r_txt.font.color.rgb = GRIS_TEXTE
        r_txt.font.name = FONT_MAIN
    doc.add_paragraph()


# ─── EN UN COUP D'ŒIL (résumé de section) ───────────────

def add_at_a_glance(doc, items, reading_time=None):
    """Encadré placé en tête d'une section, résumant son contenu en 3-5 points
    avec un temps de lecture estimé. `items` est une liste de chaînes.
    """
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F2FE")
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="6A6AF4"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="6A6AF4"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="6A6AF4"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="6A6AF4"/>'
        f'</w:tcBorders>'
    ))
    # En-tête
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    r_title = p.add_run("👁️  EN UN COUP D'ŒIL")
    r_title.bold = True
    r_title.font.size = Pt(10)
    r_title.font.color.rgb = BLEU_FONCE
    r_title.font.name = FONT_MAIN
    if reading_time:
        r_time = p.add_run(f"   \u2022   \u23f1 {reading_time} de lecture")
        r_time.italic = True
        r_time.font.size = Pt(9)
        r_time.font.color.rgb = GRIS_SEC
        r_time.font.name = FONT_MAIN

    # Items
    for item in items:
        ip = cell.add_paragraph()
        ip.paragraph_format.space_before = Pt(1)
        ip.paragraph_format.space_after = Pt(1)
        ip.paragraph_format.left_indent = Cm(0.3)
        b = ip.add_run("•  ")
        b.bold = True
        b.font.size = Pt(10.5)
        b.font.color.rgb = BLEU_ACCENT
        b.font.name = FONT_MAIN
        t = ip.add_run(item)
        t.font.size = Pt(10.5)
        t.font.color.rgb = GRIS_TEXTE
        t.font.name = FONT_MAIN
    doc.add_paragraph()


# ─── BANDEAU « DEUX PARCOURS » ─────────────────────────

def add_parcours_block(doc,
                       left_icon, left_title, left_desc, left_target,
                       right_icon, right_title, right_desc, right_target,
                       left_bookmark=None, right_bookmark=None):
    """Bandeau à deux colonnes proposant deux parcours de lecture.
    Si `left_bookmark` / `right_bookmark` sont fournis, le libellé cible
    correspondant devient un hyperlien interne cliquable.
    """
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for i, (icon, title, desc, target, bookmark) in enumerate([
        (left_icon, left_title, left_desc, left_target, left_bookmark),
        (right_icon, right_title, right_desc, right_target, right_bookmark),
    ]):
        cell = table.cell(0, i)
        cell.width = Cm(8)
        set_cell_shading(cell, "E3E3FD" if i == 0 else "E8F1FA")
        border_color = "000091" if i == 0 else "0063CB"
        cell._tc.get_or_add_tcPr().append(parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="{border_color}"/>'
            f'  <w:left w:val="single" w:sz="6" w:space="0" w:color="{border_color}"/>'
            f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{border_color}"/>'
            f'  <w:right w:val="single" w:sz="6" w:space="0" w:color="{border_color}"/>'
            f'</w:tcBorders>'
        ))
        # Icône
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(icon)
        r.font.size = Pt(24)
        # Titre
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(2)
        r = p2.add_run(title)
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = BLEU_FONCE if i == 0 else BLEU_NOTE_TEXTE
        r.font.name = FONT_MAIN
        # Description
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after = Pt(3)
        r = p3.add_run(desc)
        r.font.size = Pt(10)
        r.font.color.rgb = GRIS_TEXTE
        r.font.name = FONT_MAIN
        # Cible
        p4 = cell.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p4.paragraph_format.space_before = Pt(0)
        p4.paragraph_format.space_after = Pt(4)
        link_color = "000091" if i == 0 else "0063CB"
        # La flèche reste en texte statique.
        r = p4.add_run("→  ")
        r.bold = True
        r.italic = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = BLEU_FONCE if i == 0 else BLEU_NOTE_TEXTE
        r.font.name = FONT_MAIN
        # Le libellé devient un hyperlien si bookmark est fourni.
        if bookmark:
            _add_internal_hyperlink(p4, target, bookmark,
                                    color_hex=link_color, bold=True, italic=True,
                                    size_pt=9.5, font_name=FONT_MAIN,
                                    underline=True)
        else:
            r2 = p4.add_run(target)
            r2.bold = True
            r2.italic = True
            r2.font.size = Pt(9.5)
            r2.font.color.rgb = BLEU_FONCE if i == 0 else BLEU_NOTE_TEXTE
            r2.font.name = FONT_MAIN
    doc.add_paragraph()


# ─── SCHÉMA DU PIPELINE (ASCII) ───────────────────────

def add_pipeline_diagram(doc):
    """Schéma ASCII du pipeline de traitement."""
    diagram = (
        "  LiDAR (.laz)  \u2500\u2500\u25b6  MNT (.tif)  \u2500\u2500\u25b6  RVT (.tif)  \u2500\u2500\u25b6  IA (ONNX)  \u2500\u2500\u25b6  Shapefiles (.gpkg)\n"
        "       \u2502                \u2502                \u2502                 \u2502                \u2502\n"
        "  mode ign_laz   mode existing_mnt   mode existing_rvt   détection par    shapefiles par\n"
        "  mode local_laz                                         classe           classe détectée"
    )
    add_code_block(doc, diagram)


# ─── LÉGENDES DE FIGURES / TABLEAUX (champs SEQ) ─────────

def add_figure_caption(doc, text):
    """Légende de figure avec numérotation automatique (champ SEQ Figure)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Figure ")
    r.italic = True; r.bold = True
    r.font.size = Pt(9.5); r.font.color.rgb = BLEU_SECONDAIRE
    r.font.name = FONT_MAIN
    seq = _insert_field(p, ' SEQ Figure \\* ARABIC ')
    seq.italic = True; seq.bold = True
    seq.font.size = Pt(9.5); seq.font.color.rgb = BLEU_SECONDAIRE
    seq.font.name = FONT_MAIN
    r2 = p.add_run(f"  \u2014  {text}")
    r2.italic = True
    r2.font.size = Pt(9.5); r2.font.color.rgb = GRIS_SEC
    r2.font.name = FONT_MAIN
    return p


def add_table_caption(doc, text):
    """Légende de tableau avec numérotation automatique (champ SEQ Tableau)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Tableau ")
    r.italic = True; r.bold = True
    r.font.size = Pt(9.5); r.font.color.rgb = BLEU_SECONDAIRE
    r.font.name = FONT_MAIN
    seq = _insert_field(p, ' SEQ Tableau \\* ARABIC ')
    seq.italic = True; seq.bold = True
    seq.font.size = Pt(9.5); seq.font.color.rgb = BLEU_SECONDAIRE
    seq.font.name = FONT_MAIN
    r2 = p.add_run(f"  \u2014  {text}")
    r2.italic = True
    r2.font.size = Pt(9.5); r2.font.color.rgb = GRIS_SEC
    r2.font.name = FONT_MAIN
    return p


# ─── TABLE DES FIGURES / TABLEAUX (champs TOC avec \\c) ──────

def add_table_of_figures(doc):
    """Insère une table des figures automatique."""
    add_heading2(doc, "Table des figures")
    p = doc.add_paragraph()
    _insert_field(p, ' TOC \\h \\z \\c "Figure" ', dirty=True)


def add_table_of_tables(doc):
    """Insère une table des tableaux automatique."""
    add_heading2(doc, "Table des tableaux")
    p = doc.add_paragraph()
    _insert_field(p, ' TOC \\h \\z \\c "Tableau" ', dirty=True)


# ─── CONSTRUCTION DU DOCUMENT ───────────────────────────────

def enable_auto_update_fields(doc):
    """Force Word à mettre à jour tous les champs (TOC, SEQ, PAGE, …) à
    l'ouverture du document. Sans cette directive, la TDM et les numéros de
    page restent sur leur dernière valeur calculée (souvent vide).

    On utilise parse_xml avec la déclaration de namespace explicite — c'est
    plus robuste que OxmlElement pour l'écriture dans settings.xml.
    """
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is not None:
        return  # déjà présent
    el = parse_xml(f'<w:updateFields {nsdecls("w")} w:val="true"/>')
    # Insérer en tête pour éviter tout problème d'ordre des éléments
    settings.insert(0, el)


def build_document():
    doc = Document()

    # ── Marges ──
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Styles natifs Word (titres + Normal) nécessaires à la TDM auto
    # et au volet Navigation de Word.
    configure_heading_styles(doc)

    # Force Word à recalculer la TDM / SEQ / STYLEREF / PAGE à l'ouverture
    enable_auto_update_fields(doc)

    # En-tête (titre) + pied de page (Page X / Y) sur toutes les pages.
    setup_page_header_footer(doc)

    # ══════════════════════════════════════════════════════
    # PAGE DE COUVERTURE
    # ══════════════════════════════════════════════════════

    add_cover_page(doc)

    # ══════════════════════════════════════════════════════
    # TABLE DES MATIÈRES AUTOMATIQUE (champ TOC cliquable)
    # ══════════════════════════════════════════════════════

    add_auto_toc(doc)
    add_page_break(doc)

    # ═══════════════════════════════════════════════════════
    # 1. INTRODUCTION
    # ═══════════════════════════════════════════════════════

    add_heading1(doc, "💡  1. Introduction", bookmark="section_introduction")

    add_at_a_glance(doc, [
        "Objectif et cadre d'emploi du plugin",
        "Parcours de lecture adapté à votre besoin",
        "Schéma du pipeline de bout en bout",
    ], reading_time="2 min")

    add_body(doc,
        "Archéolog'IA est un plugin QGIS qui automatise le traitement de données "
        "LiDAR aéroporté (nuages de points LAZ) et la détection d'entités "
        "archéologiques par intelligence artificielle."
    )

    add_body(doc,
        "À partir d'une zone géographique et de données LiDAR, le plugin "
        "effectue automatiquement :"
    )

    add_bullet(doc, "Le téléchargement ou l'import des données LiDAR (nuages de points)")
    add_bullet(doc, "Le calcul du Modèle Numérique de Terrain (MNT)")
    add_bullet(doc, "Le calcul des indices de visualisation (Hillshade, SVF, Local Dominance…)")
    add_bullet(doc, "La détection automatique d'entités archéologiques par vision artificielle (IA)")
    add_bullet(doc, "L'export des résultats sous forme de rasters géoréférencés, de shapefiles et d'un projet QGIS prêt à l'analyse")

    add_conseil(doc,
        "Si vous débutez, suivez d'abord le Démarrage rapide (section 3) "
        "avant d'explorer les paramètres avancés."
    )

    # Schéma du pipeline
    add_heading2(doc, "Comment ça marche — le pipeline de traitement")
    add_body(doc,
        "Le plugin enchaîne automatiquement les étapes suivantes. Selon le mode "
        "choisi, l'entrée dans le pipeline se fait à différents endroits."
    )
    add_pipeline_diagram(doc)
    add_figure_caption(doc, "Vue d'ensemble du pipeline de traitement")

    # Deux parcours de lecture
    add_heading2(doc, "Deux parcours de lecture")
    add_parcours_block(doc,
        "🚀", "Je veux démarrer vite",
        "Un scénario complet en 8 étapes avec un cas concret (mode ign_laz).",
        "Section 3 — Démarrage rapide",
        "⚙️", "Je veux tout paramétrer",
        "La configuration complète et tous les paramètres avancés.",
        "Section 5 + Annexes techniques",
        left_bookmark="section_demarrage_rapide",
        right_bookmark="section_configuration",
    )

    # ═══════════════════════════════════════════════════════
    # 2. PRÉREQUIS ET INSTALLATION
    # ═══════════════════════════════════════════════════════

    add_heading1(doc, "📋  2. Prérequis et installation")

    add_at_a_glance(doc, [
        "Logiciels requis (QGIS, PDAL)",
        "Installation du plugin en 4 étapes",
        "Structure des données à déposer dans data/",
        "Vérification automatique au premier lancement",
    ], reading_time="5 min")

    # 2.1
    add_heading2(doc, "2.1 Logiciels requis")

    add_styled_table(doc,
        ["Logiciel", "Version minimale", "Remarque"],
        [
            ["QGIS", "3.28 LTR", "Installation via OSGeo4W recommandée"],
            ["PDAL", "inclus dans QGIS OSGeo4W", "Traitement des nuages de points"],
        ],
        col_widths=[4, 4, 7]
    )

    # 2.2
    add_heading2(doc, "2.2 Installation du plugin")

    add_bullet(doc, "Télécharger l'archive .zip du plugin")
    add_bullet(doc, "Dans QGIS : Extensions → Installer/Gérer les extensions → Installer depuis un ZIP")
    add_bullet(doc, "Sélectionner le fichier .zip téléchargé")
    add_bullet(doc, "Cliquer sur Installer le plugin")

    add_body(doc,
        "Le plugin apparaît ensuite dans le menu Extensions → Archéolog'IA."
    )

    # 2.3
    add_heading2(doc, "2.3 Données requises")

    add_body(doc,
        "Le plugin nécessite trois dossiers de données à placer dans le "
        "répertoire data/ du plugin."
    )

    add_body(doc,
        "Le chemin complet du dossier data/ est affiché dans la fenêtre "
        "d'erreur si un fichier est manquant."
    )

    add_heading3(doc, "Dossier data/models/")
    add_body(doc,
        "Contient les modèles d'intelligence artificielle (format ONNX). "
        "Chaque modèle est un sous-dossier contenant :"
    )
    add_bullet(doc, "weights/best.onnx — le modèle entraîné")
    add_bullet(doc, "args.yaml — la configuration du modèle (classes, paramètres SAHI, clustering…)")
    add_bullet(doc, "classes.txt — la liste des classes détectables")

    add_screenshot_placeholder(doc, "Arborescence d'un modèle")

    add_heading3(doc, "Dossier data/third_party/cv_runner_onnx/")
    add_body(doc,
        "Contient le moteur d'inférence compilé (cv_runner_onnx.exe sur Windows). "
        "Ce fichier est fourni séparément car il dépend de l'architecture matérielle."
    )

    add_heading3(doc, "Dossier data/quadrillage_france/")
    add_body(doc,
        "Contient le fichier de la grille des dalles LiDAR IGN "
        "(TA_diff_pkk_lidarhd_classe.shp). Nécessaire uniquement pour le mode ign_laz. "
        "Ce fichier est fourni par l'IGN (~200 Mo)."
    )

    # 2.4
    add_heading2(doc, "2.4 Vérification de l'installation")
    add_body(doc,
        "Au premier lancement du plugin, une vérification automatique est effectuée. "
        "En cas d'élément manquant, un message d'erreur indique précisément ce qui est absent."
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════
    # 3. DÉMARRAGE RAPIDE
    # ═══════════════════════════════════════════════════════

    add_heading1(doc, "🚀  3. Démarrage rapide", bookmark="section_demarrage_rapide")

    add_at_a_glance(doc, [
        "Un scénario complet, du chargement à l'analyse",
        "8 étapes illustrées avec des captures",
        "Base solide pour ensuite creuser la configuration",
    ], reading_time="10 min")

    add_body(doc,
        "Ce guide illustre un cas d'usage typique : détecter des entités "
        "archéologiques sur une zone à partir du LiDAR IGN."
    )

    # Étape 1
    add_step_callout(doc, 1, "Ouvrir le plugin")
    add_body(doc, "Dans QGIS, cliquer sur Extensions → Archéolog'IA.")
    add_body(doc,
        "La fenêtre principale s'ouvre. Elle est organisée en deux parties :"
    )
    add_bullet(doc, "Partie haute : la configuration (paramètres du traitement)")
    add_bullet(doc, "Partie basse : la console de logs (suivi en temps réel)")
    add_screenshot_placeholder(doc, "Fenêtre principale en mode Simple")

    # Étape 2
    add_step_callout(doc, 2, "Choisir le mode de traitement")
    add_body(doc,
        "Dans la liste déroulante en haut à gauche, sélectionner Simple pour commencer."
    )
    add_body(doc,
        "Le mode Simple masque les paramètres avancés et utilise les valeurs par défaut optimisées."
    )

    # Étape 3
    add_step_callout(doc, 3, "Sélectionner la source de données")
    add_body(doc, "Dans la section Sources, choisir :")
    add_bullet(doc, "Mode : ign_laz")
    add_bullet(doc, "Zone d'étude : sélectionner un fichier vectoriel (.shp, .geojson) ou un fichier .txt de dalles")
    add_body(doc,
        "Le fichier vectoriel doit délimiter la zone à traiter (polygone en n'importe "
        "quelle projection, le plugin se charge de la reprojection)."
    )
    add_screenshot_placeholder(doc, "Section Sources remplie")

    # Étape 4
    add_step_callout(doc, 4, "Choisir le dossier de sortie")
    add_body(doc,
        "Cliquer sur le bouton de sélection du Dossier de sortie et choisir un dossier "
        "vide sur un disque avec suffisamment d'espace (prévoir ~5 Go par dalle LiDAR)."
    )

    # Étape 5
    add_step_callout(doc, 5, "Sélectionner les produits à générer")
    add_body(doc,
        "Dans la section MNT et Indices de visualisation, cocher les produits souhaités :"
    )

    add_styled_table(doc,
        ["Produit", "Description", "Recommandé"],
        [
            ["MNT", "Modèle Numérique de Terrain (raster altitude)", "✅ Toujours"],
            ["M-HS", "Hillshade multi-directionnel", "✅ Recommandé"],
            ["SVF", "Sky View Factor", "✅ Recommandé"],
            ["LD", "Local Dominance", "✅ Pour la CV"],
            ["SLO", "Pente (Slope)", "Optionnel"],
            ["SLRM", "Simple Local Relief Model", "Optionnel"],
            ["VAT", "Visualisation Archéologique Totale", "Optionnel"],
            ["DENSITE", "Carte de densité de points", "Optionnel"],
        ],
        col_widths=[3, 7, 4]
    )

    # Étape 6
    add_step_callout(doc, 6, "Activer la Computer Vision (optionnel)")
    add_body(doc, "Dans la section Détection automatique (Computer Vision) :")
    add_bullet(doc, "Cocher Activer la Computer Vision")
    add_bullet(doc, "Sélectionner un modèle dans la liste déroulante")
    add_bullet(doc, "Vérifier que les classes à détecter sont bien cochées")
    add_screenshot_placeholder(doc, "Section CV configurée")

    # Étape 7
    add_step_callout(doc, 7, "Lancer le traitement")
    add_body(doc, "Cliquer sur le bouton ▶ Lancer en bas de la fenêtre.")
    add_body(doc,
        "Le traitement démarre. La progression est visible dans la barre de statut "
        "et la console de logs."
    )
    add_conseil(doc,
        "Le temps de traitement dépend du nombre de dalles et des options choisies. "
        "Prévoir environ 5 à 15 minutes par dalle LiDAR (~1 km²) sur un poste standard."
    )

    # Étape 8
    add_step_callout(doc, 8, "Consulter les résultats")
    add_body(doc,
        "À la fin du traitement, les couches sont automatiquement chargées dans QGIS :"
    )
    add_bullet(doc, "Les rasters (MNT, indices RVT) dans le groupe Indices")
    add_bullet(doc, "Les shapefiles de détection dans le groupe Détections")
    add_bullet(doc, "Le projet QGIS consolidé est enregistré dans le dossier de sortie")
    add_screenshot_placeholder(doc, "Résultats chargés dans QGIS")

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════
    # 4. LES 4 MODES DE TRAITEMENT
    # ═══════════════════════════════════════════════════════

    add_heading1(doc, "🗺️  4. Les 4 modes de traitement")

    add_at_a_glance(doc, [
        "Tableau comparatif en un coup d'œil",
        "Quand utiliser chaque mode",
        "Détail et format d'entrée pour chacun",
    ], reading_time="6 min")

    add_body(doc,
        "Le plugin prend en charge 4 modes selon le type de données disponibles."
    )

    # Tableau comparatif des modes
    add_heading2(doc, "4.0 Synthèse des 4 modes")
    add_styled_table(doc,
        ["Mode", "Quand l'utiliser", "Entrée", "Internet", "Temps/dalle"],
        [
            ["ign_laz", "France, pas de LAZ local", "Fichier vectoriel ou .txt", "✅ Oui", "10–15 min"],
            ["local_laz", "LAZ déjà téléchargés", "Dossier .laz / .las", "❌ Non", "8–12 min"],
            ["existing_mnt", "MNT déjà produit", "Dossier .tif / .asc", "❌ Non", "3–5 min"],
            ["existing_rvt", "Indices RVT prêts", "Dossier .tif", "❌ Non", "1–3 min"],
        ],
        col_widths=[3, 4.5, 4.5, 2, 3]
    )
    add_table_caption(doc, "Comparaison des 4 modes de traitement")

    # 4.1
    add_heading2(doc, "4.1 Mode ign_laz — Données LiDAR IGN (recommandé)")

    add_body_rich(doc, [
        ("Quand l'utiliser : ", True, False, GRIS_TEXTE),
        ("vous souhaitez traiter une zone en France métropolitaine sans disposer "
         "des dalles LiDAR en local.", False, False, GRIS_TEXTE),
    ])

    add_body_rich(doc, [("Fonctionnement :", True, False, GRIS_TEXTE)])
    add_bullet(doc, "Le plugin interroge la grille nationale IGN pour identifier les dalles couvrant la zone")
    add_bullet(doc, "Les dalles sont téléchargées automatiquement depuis les serveurs IGN")
    add_bullet(doc, "Les traitements sont appliqués")

    add_body_rich(doc, [("Entrée acceptée :", True, False, GRIS_TEXTE)])
    add_bullet(doc, "Un fichier vectoriel (.shp, .geojson, .gpkg) délimitant la zone d'étude")
    add_bullet(doc, "Un fichier texte (.txt) listant les URLs de dalles IGN (une par ligne, format nom_dalle,url)")

    add_avertissement(doc, "Connexion internet requise pour le téléchargement des dalles.")

    # 4.2
    add_heading2(doc, "4.2 Mode local_laz — Fichiers LAZ locaux")

    add_body_rich(doc, [
        ("Quand l'utiliser : ", True, False, GRIS_TEXTE),
        ("vous disposez déjà des fichiers LAZ sur votre poste ou sur un serveur local.", False, False, GRIS_TEXTE),
    ])

    add_body(doc,
        "Fonctionnement : le plugin scanne le dossier spécifié, fusionne les dalles "
        "adjacentes si nécessaire, et applique les mêmes traitements que le mode IGN."
    )

    add_body_rich(doc, [
        ("Entrée acceptée : ", True, False, GRIS_TEXTE),
        ("un dossier contenant les fichiers .laz ou .las", False, False, GRIS_TEXTE),
    ])

    # 4.3
    add_heading2(doc, "4.3 Mode existing_mnt — MNT déjà calculés")

    add_body_rich(doc, [
        ("Quand l'utiliser : ", True, False, GRIS_TEXTE),
        ("vous disposez déjà de MNT (rasters d'altitude) mais pas encore des "
         "indices de visualisation.", False, False, GRIS_TEXTE),
    ])

    add_body(doc,
        "Fonctionnement : le plugin charge les fichiers MNT (.tif ou .asc), calcule "
        "les indices RVT sélectionnés, puis exécute éventuellement la Computer Vision."
    )

    add_body_rich(doc, [
        ("Entrée acceptée : ", True, False, GRIS_TEXTE),
        ("un dossier contenant des fichiers MNT (.tif, .tiff, .asc)", False, False, GRIS_TEXTE),
    ])

    # 4.4
    add_heading2(doc, "4.4 Mode existing_rvt — Indices de visualisation déjà produits")

    add_body_rich(doc, [
        ("Quand l'utiliser : ", True, False, GRIS_TEXTE),
        ("vous disposez déjà des indices de visualisation (rasters RVT) et souhaitez "
         "uniquement lancer la détection automatique par IA.", False, False, GRIS_TEXTE),
    ])

    add_body(doc,
        "Fonctionnement : le plugin copie et normalise les fichiers TIF, puis lance "
        "la Computer Vision."
    )

    add_body_rich(doc, [
        ("Entrée acceptée : ", True, False, GRIS_TEXTE),
        ("un dossier contenant des fichiers .tif / .tiff", False, False, GRIS_TEXTE),
    ])

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════
    # 5. CONFIGURATION DE BASE
    # ═══════════════════════════════════════════════════════

    add_heading1(doc, "⚙️  5. Configuration de base", bookmark="section_configuration")

    add_at_a_glance(doc, [
        "Sources, produits RVT et Computer Vision",
        "Sauvegarde / chargement de configuration en JSON",
        "Différence Simple / Expert",
    ], reading_time="5 min")

    # 5.1
    add_heading2(doc, "5.1 Section Sources")

    add_styled_table(doc,
        ["Paramètre", "Description"],
        [
            ["Mode", "Le mode de traitement (voir section 4)"],
            ["Zone d'étude / Dossier source", "Fichier vectoriel, fichier .txt ou dossier selon le mode"],
            ["Dossier de sortie", "Dossier où seront écrits tous les fichiers produits"],
        ],
        col_widths=[5, 10]
    )

    # 5.2
    add_heading2(doc, "5.2 Section MNT et Indices de visualisation")

    add_body(doc,
        "Cocher les produits à calculer. Chaque produit coché sera calculé pour "
        "chaque dalle traitée."
    )

    add_body(doc,
        "En mode Simple, les paramètres de calcul utilisent les valeurs par défaut. "
        "En mode Expert, chaque indice expose ses paramètres détaillés (voir Annexe A)."
    )

    add_screenshot_placeholder(doc, "Section produits")

    # 5.3
    add_heading2(doc, "5.3 Sauvegarde et chargement de configuration")

    add_body(doc,
        "Les boutons Sauvegarder config et Charger config (en haut de la fenêtre) "
        "permettent d'enregistrer et de recharger un profil de configuration complet "
        "au format .json."
    )

    add_body(doc, "Cela permet de :")
    add_bullet(doc, "Reprendre un traitement avec les mêmes paramètres")
    add_bullet(doc, "Partager une configuration entre utilisateurs")
    add_bullet(doc, "Archiver les paramètres d'un traitement passé")

    # 5.4
    add_heading2(doc, "5.4 Mode Simple / Expert")

    add_body(doc,
        "La liste déroulante Mode en haut de la fenêtre bascule entre deux niveaux "
        "de configuration :"
    )
    add_bullet(doc, "Simple : seuls les paramètres essentiels sont visibles")
    add_bullet(doc, "Expert : tous les paramètres sont accessibles (résolutions, paramètres RVT, options CV avancées, performances)")

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════
    # 6. DÉTECTION AUTOMATIQUE PAR COMPUTER VISION
    # ═══════════════════════════════════════════════════════

    add_heading1(doc, "🎯  6. Détection automatique par Computer Vision")

    add_at_a_glance(doc, [
        "Principe de la détection automatique par IA",
        "Configurer plusieurs runs en parallèle",
        "Sélectionner les classes à détecter",
        "Paramètres clés : seuils, SAHI",
    ], reading_time="8 min")

    add_body(doc,
        "La Computer Vision (CV) utilise des modèles d'intelligence artificielle "
        "pour détecter automatiquement des entités archéologiques dans les indices "
        "de visualisation."
    )

    # 6.1
    add_heading2(doc, "6.1 Activation")
    add_body(doc,
        "Dans la section Détection automatique, cocher Activer la Computer Vision."
    )

    # 6.2
    add_heading2(doc, "6.2 Configuration des runs")
    add_body(doc,
        "Le plugin permet de configurer plusieurs runs de détection en parallèle. "
        "Chaque run associe :"
    )
    add_bullet(doc, "Un modèle IA")
    add_bullet(doc, "Un indice RVT cible (le raster sur lequel le modèle s'applique)")

    add_body(doc, "Pour ajouter un run : cliquer sur + Ajouter un run")
    add_body(doc, "Pour supprimer un run : sélectionner la ligne et cliquer sur Supprimer")

    add_screenshot_placeholder(doc, "Tableau des runs CV")

    # 6.3
    add_heading2(doc, "6.3 Sélection des classes")
    add_body(doc,
        "Pour chaque run, une liste des classes détectables par le modèle s'affiche. "
        "Décocher une classe l'exclut des résultats."
    )
    add_conseil(doc,
        "Si toutes les classes sont décochées pour un run, l'inférence est "
        "complètement ignorée pour ce run (gain de temps)."
    )

    # 6.4
    add_heading2(doc, "6.4 Paramètres principaux")

    add_styled_table(doc,
        ["Paramètre", "Description", "Valeur par défaut"],
        [
            ["Seuil de confiance", "Score minimum pour qu'une détection soit conservée (0–1)", "0.3"],
            ["Seuil IOU", "Seuil de recouvrement pour la suppression des doublons", "0.5"],
            ["Images annotées", "Générer des PNG avec les détections superposées", "Non"],
            ["Shapefiles", "Générer les shapefiles de détection", "Oui"],
        ],
        col_widths=[4, 7, 3.5]
    )

    # 6.5
    add_heading2(doc, "6.5 Résultats de la CV")
    add_body(doc,
        "À la fin du traitement, les résultats de détection sont disponibles sous forme :"
    )
    add_bullet(doc,
        "Shapefiles par classe : un fichier .shp par classe détectée, avec les "
        "attributs confidence, class_name, area_m2"
    )
    add_bullet(doc,
        "Projet QGIS consolidé : detections_validation.qgs à la racine du dossier "
        "de sortie, avec toutes les couches de détection chargées et stylées avec "
        "une couleur distincte par classe"
    )
    add_bullet(doc,
        "Images annotées (optionnel) : PNG avec les détections dessinées en surimpression"
    )

    # 6.6
    add_heading2(doc, "6.6 Modèles fournis par défaut")
    add_body(doc,
        "Trois modèles pré-entraînés sont fournis dans data/models/. Chacun "
        "s'applique à un indice RVT précis et détecte des classes spécifiques."
    )
    add_styled_table(doc,
        ["Modèle", "Indice cible", "Classes détectées", "Particularité"],
        [
            [
                "formes_lineaires_ld_a15_rmin10_rm_rfdetr_seg_1",
                "LD (A15, Rmin10)",
                "chemin_creux, parcellaire, talus-fosse_fossebutte",
                "RF-DETR-Seg, segmentation d'instances",
            ],
            [
                "run_rf_detr_1",
                "LD",
                "charbonnière, circular_depression, four",
                "Détection de structures ponctuelles",
            ],
            [
                "verdun_3_classes_1",
                "SVF",
                "abri, cratere_obus, tranchees_et_boyaux",
                "Clustering des cratères → zone_crateres",
            ],
        ],
        col_widths=[5.5, 2.5, 5, 4]
    )
    add_table_caption(doc, "Modèles fournis avec le plugin")

    add_note(doc,
        "Lorsque vous ajoutez un run et sélectionnez un modèle, le plugin sélectionne "
        "automatiquement le bon indice RVT et l'aire minimale de détection à partir "
        "du fichier training_params.json du modèle."
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════
    # 7. COMPRENDRE LES SORTIES
    # ═══════════════════════════════════════════════════════

    add_heading1(doc, "📦  7. Comprendre les sorties")

    add_at_a_glance(doc, [
        "Arborescence complète du dossier de sortie",
        "Le projet QGIS consolidé (point d'entrée recommandé)",
        "Attributs des shapefiles de détection",
        "Logs et métadonnées",
    ], reading_time="7 min")

    # 7.1
    add_heading2(doc, "7.1 Arborescence du dossier de sortie")

    add_body(doc,
        "Après un traitement complet, le dossier de sortie contient les éléments "
        "suivants (exemple concret tiré d'un vrai run sur 2 dalles IGN) :"
    )
    add_code_block(doc, """📁 dossier_sortie/
├── 📁 sources/
│   └── 📁 dalles/
│       ├── 📄 LHD_FXX_0822_6329_PTS_O_LAMB93_IGN69.copc.laz
│       └── 📄 LHD_FXX_0823_6329_PTS_O_LAMB93_IGN69.copc.laz
├── 📁 intermediaires/                 # Fichiers techniques (MNT, LAZ fusionnés)
│   ├── 📄 *_MNT_A_0M50_LAMB93_IGN69.tif
│   ├── 📄 *_LD_A15_Rmin10_Rmax20_H1p7_V1_A_LAMB93.tif
│   └── 📄 *_PTS_O_LAMB93_IGN69_merged.laz
├── 📁 indices/                        # Rasters utilisables en SIG
│   ├── 📁 MNT/tif/                    # MNT géoréférencés + index.vrt (mosaïque)
│   ├── 📁 LD/
│   │   ├── 📁 tif/                    # Local Dominance GeoTIFF
│   │   └── 📁 png/                    # PNG 8-bit (entrée Computer Vision)
│   └── 📁 M_HS/ SVF/ SLO/ ...         # Autres indices RVT (si activés)
├── 📁 detections/                     # Résultats Computer Vision
│   ├── 📄 detections_validation.qgs   # ⭐ Projet QGIS consolidé (point d'entrée)
│   └── 📁 <nom_du_modèle>/
│       ├── 📁 shapefiles/             # 📄 detections_<INDICE>.gpkg (toutes classes)
│       ├── 📁 raw_detections/         # Sorties brutes (JSON/TXT + classes.txt)
│       └── 📁 annotated_images/       # PNG annotés + legend.png + index.vrt
├── 📄 dalles_urls.txt                 # Liste des dalles IGN traitées
├── 📄 fichier_tri.txt                 # Trace du tri initial
├── 📄 metadata.json                   # Configuration complète du run (rechargeable)
└── 📄 pipeline_log_AAAAMMJJ_HHMMSS.txt  # Journal complet du traitement""")
    add_figure_caption(doc, "Arborescence complète d'un dossier de sortie (exemple 2 dalles)")

    add_note(doc,
        "Le projet QGIS consolidé (detections_validation.qgs) est placé à l'intérieur "
        "du dossier detections/ — c'est le point d'entrée recommandé pour valider et "
        "analyser les résultats."
    )

    # Nommage des fichiers raster
    add_heading3(doc, "Convention de nommage des fichiers raster")
    add_body(doc, "Les fichiers suivent la convention de nommage IGN LiDAR HD :")
    add_code_block(doc,
        "LHD_FXX_<X_km>_<Y_km>_<PRODUIT>_A_LAMB93.tif\n\n"
        "Exemple : LHD_FXX_0822_6329_LD_A15_Rmin10_Rmax20_H1p7_V1_A_LAMB93.tif"
    )

    # 7.2
    add_heading2(doc, "7.2 Le projet QGIS consolidé")
    add_body(doc,
        "Le fichier detections_validation.qgs est automatiquement chargé dans QGIS "
        "à la fin du traitement. Il contient :"
    )
    add_bullet(doc, "Toutes les couches de shapefiles de détection")
    add_bullet(doc, "Une couleur distincte par classe de détection")
    add_bullet(doc, "Les rasters VRT (mosaïques) pour chaque indice RVT")

    add_body(doc,
        "C'est le point d'entrée recommandé pour l'analyse des résultats."
    )

    # 7.3
    add_heading2(doc, "7.3 Les shapefiles de détection")
    add_body(doc, "Chaque shapefile de détection contient les attributs suivants :")

    add_styled_table(doc,
        ["Attribut", "Description"],
        [
            ["class_name", "Nom de la classe détectée"],
            ["confidence", "Score de confiance du modèle (0–1)"],
            ["area_m2", "Surface de la détection en m²"],
            ["cluster_id", "Identifiant de groupe (si clustering activé)"],
        ],
        col_widths=[4, 11]
    )

    add_body(doc, "Pour les classes de clustering (ex : zone_crateres) :")

    add_styled_table(doc,
        ["Attribut", "Description"],
        [
            ["nb_detect", "Nombre de détections individuelles dans le groupe"],
            ["mean_confidence", "Confiance moyenne des détections du groupe"],
            ["density", "Densité de détections par hectare"],
        ],
        col_widths=[4, 11]
    )

    # 7.4
    add_heading2(doc, "7.4 Le fichier metadata.json")
    add_body(doc,
        "Ce fichier JSON documente les paramètres du traitement effectué : date, mode, "
        "nombre de dalles, modèles utilisés, produits générés. Il sert de trace pour "
        "l'archivage et la reproductibilité."
    )

    # 7.5
    add_heading2(doc, "7.5 Les logs")
    add_body(doc,
        "Le fichier pipeline_log_AAAAMMJJ_HHMMSS.txt contient le journal complet du "
        "traitement. Il est utile pour diagnostiquer un problème ou vérifier le "
        "déroulement du pipeline."
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════
    # 8. PERFORMANCE & MATÉRIEL
    # ══════════════════════════════════════════════════════

    add_heading1(doc, "⚡  8. Performance & matériel")

    add_at_a_glance(doc, [
        "Configuration matérielle recommandée (RAM, CPU, GPU)",
        "Temps typiques par dalle et par étape du pipeline",
        "Leviers d'optimisation (Workers, résolution MNT…)",
        "Astuces pour traiter de grandes zones",
    ], reading_time="4 min")

    add_heading2(doc, "8.1 Configuration recommandée")
    add_styled_table(doc,
        ["Ressource", "Minimum", "Recommandé"],
        [
            ["RAM", "16 Go", "32 Go ou plus"],
            ["CPU", "4 cœurs (x86_64)", "8 cœurs physiques ou plus"],
            ["GPU", "Optionnel (CPU OK)", "NVIDIA RTX (accélère la CV ×3 à ×10)"],
            ["Disque", "SSD, 50 Go libres", "NVMe, 200 Go+ pour grandes zones"],
        ],
        col_widths=[4, 5.5, 5.5]
    )
    add_table_caption(doc, "Configuration matérielle")

    add_heading2(doc, "8.2 Temps de traitement indicatifs")
    add_body(doc,
        "Ordres de grandeur mesurés sur une machine 8 cœurs / 32 Go RAM, pour une "
        "dalle IGN standard (1 km² ≈ 200 millions de points)."
    )
    add_styled_table(doc,
        ["Étape", "Temps par dalle"],
        [
            ["Téléchargement IGN (mode ign_laz)", "1–3 min (selon la connexion)"],
            ["Calcul du MNT (PDAL)", "2–5 min"],
            ["Calcul d'un indice RVT (hors MNT)", "30 s – 2 min"],
            ["Inférence CV (1 modèle, CPU)", "3–8 min"],
            ["Inférence CV (1 modèle, GPU NVIDIA)", "30 s – 1 min"],
            ["Post-traitement (fusion, clustering)", "< 30 s"],
        ],
        col_widths=[9, 6]
    )
    add_table_caption(doc, "Temps indicatifs par étape du pipeline")

    add_heading2(doc, "8.3 Leviers d'optimisation")
    add_body(doc, "En mode Expert, plusieurs leviers permettent d'accélérer un traitement :")
    add_bullet(doc, "Augmenter le nombre de Workers (jusqu'au nombre de cœurs physiques)")
    add_bullet(doc, "Baisser la résolution du MNT (ex : 1 m au lieu de 0,5 m → divise le temps par 4)")
    add_bullet(doc, "Décocher les indices RVT non exploités par la Computer Vision")
    add_bullet(doc, "Désactiver la génération d'images annotées si non nécessaire")
    add_bullet(doc, "Traiter la zone en sous-zones successives plutôt qu'en une seule passe")

    add_avertissement(doc,
        "Plus de Workers = plus de RAM consommée. Sur un poste 16 Go, "
        "rester sur 2–4 Workers pour éviter les erreurs mémoire."
    )

    add_heading2(doc, "8.4 Pour les grandes zones")
    add_succes(doc,
        "Pour une zone > 50 km² : commencez par un test sur 1–2 dalles pour valider "
        "votre configuration, puis lancez la zone complète avec le profil sauvegardé."
    )
    add_example(doc,
        "Une zone de 100 km² avec MNT + LD + SVF + 1 modèle CV, sur un poste 8 cœurs / "
        "32 Go / GPU RTX : prévoir environ 3 à 5 heures de traitement."
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════
    # 9. DÉPANNAGE
    # ══════════════════════════════════════════════════════

    add_heading1(doc, "🩹  9. Dépannage")

    add_at_a_glance(doc, [
        "Erreurs courantes sous forme de cartes Q/R",
        "Problèmes de lancement du plugin",
        "Computer Vision qui ne produit pas de détection",
    ], reading_time="3 min")

    depannage = [
        (
            "Le plugin ne se lance pas",
            [
                "Vérifier que QGIS est bien installé via le paquet OSGeo4W.",
                "Relancer QGIS en tant qu'administrateur si un problème de permissions est affiché.",
            ]
        ),
        (
            '"Fichier quadrillage IGN introuvable"',
            [
                "Le fichier data/quadrillage_france/TA_diff_pkk_lidarhd_classe.shp est absent.",
                "Télécharger le fichier de grille IGN et le placer dans ce dossier.",
            ]
        ),
        (
            '"Runner ONNX introuvable"',
            [
                "Le fichier data/third_party/cv_runner_onnx/windows/cv_runner_onnx.exe est absent.",
                "Placer le fichier cv_runner_onnx.exe fourni séparément dans ce dossier.",
            ]
        ),
        (
            '"Aucun modèle disponible" dans la liste CV',
            [
                "Le dossier data/models/ est vide ou mal structuré.",
                "Chaque modèle doit être dans un sous-dossier contenant au minimum weights/best.onnx.",
            ]
        ),
        (
            "Le traitement s'arrête en cours de route",
            [
                "Consulter le fichier de log (pipeline_log_*.txt) dans le dossier de sortie.",
                "Les erreurs sont préfixées par [ERROR].",
            ]
        ),
        (
            "La détection CV ne produit aucun résultat",
            [
                "Vérifier que le seuil de confiance n'est pas trop élevé (essayer 0.1 pour commencer).",
                "Vérifier que l'indice cible du modèle correspond bien à un indice calculé (ex : un modèle entraîné sur LD nécessite que LD soit coché).",
                "Vérifier que des classes sont bien sélectionnées dans la liste.",
            ]
        ),
        (
            "Les couches ne se chargent pas dans QGIS après le traitement",
            [
                "Les résultats sont dans le dossier de sortie et peuvent être chargés manuellement via Couche → Ajouter une couche.",
                "Ouvrir le fichier detections_validation.qgs directement depuis QGIS.",
            ]
        ),
    ]

    for question, solutions in depannage:
        add_depannage_card(doc, question, solutions)

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════
    # ANNEXES
    # ═══════════════════════════════════════════════════════

    add_heading1(doc, "Annexes techniques")

    # ── ANNEXE A ──
    add_heading1(doc, "Annexe A — Paramètres des indices de visualisation (RVT)")

    add_body(doc,
        "Ces paramètres sont accessibles en mode Expert, dans la section "
        "Paramètres des indices de visualisation."
    )

    # A.1
    add_heading2(doc, "A.1 Hillshade multi-directionnel (M-HS)")
    add_body(doc,
        "Le Hillshade simule l'éclairage du terrain par le soleil, révélant les micro-reliefs."
    )

    add_styled_table(doc,
        ["Paramètre", "Description", "Valeur par défaut"],
        [
            ["num_directions", "Nombre de directions d'éclairage", "16"],
            ["sun_elevation", "Angle d'élévation solaire en degrés (0–90)", "35"],
            ["ve_factor", "Facteur d'exagération verticale", "1"],
            ["save_as_8bit", "Sauvegarder en 8 bits (plus léger)", "Oui"],
        ],
        col_widths=[4, 7, 3.5]
    )

    add_conseil(doc,
        "Un sun_elevation bas (15–35°) accentue les ombres et révèle mieux les micro-reliefs."
    )

    # A.2
    add_heading2(doc, "A.2 Sky View Factor (SVF)")
    add_body(doc,
        "Le SVF mesure la proportion de ciel visible depuis chaque point du terrain. "
        "Les creux apparaissent sombres, les bosses apparaissent clairs."
    )

    add_styled_table(doc,
        ["Paramètre", "Description", "Valeur par défaut"],
        [
            ["num_directions", "Nombre de directions de calcul", "16"],
            ["radius", "Rayon de recherche en pixels", "10"],
            ["noise_remove", "Suppression du bruit (0 = désactivé)", "0"],
            ["ve_factor", "Facteur d'exagération verticale", "1"],
            ["save_as_8bit", "Sauvegarder en 8 bits", "Oui"],
        ],
        col_widths=[4, 7, 3.5]
    )

    add_conseil(doc,
        "Augmenter le radius améliore la détection de grandes structures mais ralentit le calcul."
    )

    # A.3
    add_heading2(doc, "A.3 Local Dominance (LD)")
    add_body(doc,
        "Le LD mesure à quel point un point domine localement son environnement. "
        "Particulièrement efficace pour la détection de structures linéaires et "
        "de reliefs en creux."
    )

    add_styled_table(doc,
        ["Paramètre", "Description", "Valeur par défaut"],
        [
            ["min_radius", "Rayon minimum en pixels", "10"],
            ["max_radius", "Rayon maximum en pixels", "20"],
            ["angular_res", "Résolution angulaire en degrés", "15"],
            ["observer_h", "Hauteur de l'observateur en mètres", "1.7"],
            ["ve_factor", "Facteur d'exagération verticale", "1"],
            ["save_as_8bit", "Sauvegarder en 8 bits", "Oui"],
        ],
        col_widths=[4, 7, 3.5]
    )

    # A.4
    add_heading2(doc, "A.4 Pente (SLO / Slope)")

    add_styled_table(doc,
        ["Paramètre", "Description", "Valeur par défaut"],
        [
            ["unit", "Unité de la pente : 0 = degrés, 1 = pourcentage", "0"],
            ["ve_factor", "Facteur d'exagération verticale", "1"],
            ["save_as_8bit", "Sauvegarder en 8 bits", "Oui"],
        ],
        col_widths=[4, 7, 3.5]
    )

    # A.5
    add_heading2(doc, "A.5 Simple Local Relief Model (SLRM)")
    add_body(doc,
        "Le SLRM soustrait une surface lissée au MNT pour isoler les micro-reliefs."
    )

    add_styled_table(doc,
        ["Paramètre", "Description", "Valeur par défaut"],
        [
            ["radius", "Rayon de lissage en pixels", "20"],
            ["ve_factor", "Facteur d'exagération verticale", "1"],
            ["save_as_8bit", "Sauvegarder en 8 bits", "Oui"],
        ],
        col_widths=[4, 7, 3.5]
    )

    # A.6
    add_heading2(doc, "A.6 Visualisation Archéologique Totale (VAT)")

    add_styled_table(doc,
        ["Paramètre", "Description", "Valeur par défaut"],
        [
            ["terrain_type", "Type de terrain : 0 = général, 1 = plat", "0"],
            ["save_as_8bit", "Sauvegarder en 8 bits", "Oui"],
        ],
        col_widths=[4, 7, 3.5]
    )

    # A.7
    add_heading2(doc, "A.7 Résolution et filtres MNT")
    add_body(doc,
        "Ces paramètres se trouvent dans la section Paramètres MNT (mode Expert)."
    )

    add_styled_table(doc,
        ["Paramètre", "Description", "Valeur par défaut"],
        [
            ["Résolution MNT", "Taille d'un pixel en mètres", "0.5 m"],
            ["Résolution densité", "Résolution de la carte de densité", "1.0 m"],
            ["Chevauchement dalle", "Marge de chevauchement entre dalles en mètres", "20 m"],
            ["Filtre PDAL", "Expression de filtrage des classes de points LiDAR", "Classification = 2 OR 6 OR 66 OR 67 OR 9"],
            ["Workers", "Nombre de processeurs parallèles", "2"],
        ],
        col_widths=[4, 7, 3.5]
    )

    add_avertissement(doc,
        "Le filtre PDAL par défaut conserve le sol (2), les bâtiments (6), et les "
        "classes spécifiques IGN LiDAR HD (66, 67, 9). Ne modifier que si vous "
        "connaissez la classification ASPRS de vos données."
    )

    add_page_break(doc)

    # ── ANNEXE B ──
    add_heading1(doc, "Annexe B — Paramètres avancés Computer Vision")

    # B.1
    add_heading2(doc, "B.1 SAHI (Slicing Aided Hyper Inference)")
    add_body(doc,
        "SAHI est une technique qui découpe l'image en tuiles se chevauchant pour "
        "améliorer la détection des petits objets."
    )

    add_styled_table(doc,
        ["Paramètre", "Description", "Valeur par défaut"],
        [
            ["slice_height", "Hauteur de chaque tuile en pixels", "450"],
            ["slice_width", "Largeur de chaque tuile en pixels", "450"],
            ["overlap_ratio", "Taux de chevauchement entre tuiles (0–1)", "0.2"],
        ],
        col_widths=[4, 7, 3.5]
    )

    add_avertissement(doc,
        "Ces valeurs sont définies dans le fichier args.yaml de chaque modèle et "
        "correspondent aux paramètres d'entraînement. Ne les modifier que si vous "
        "savez ce que vous faites."
    )

    # B.2
    add_heading2(doc, "B.2 Autres paramètres CV (mode Expert)")

    add_styled_table(doc,
        ["Paramètre", "Description", "Valeur par défaut"],
        [
            ["force_reprocess", "Ré-inférer même si les labels existent déjà", "Non"],
            ["scan_all", "Traiter toutes les images (sinon seulement la première)", "Oui"],
            ["generate_annotated_images", "Générer des PNG avec détections superposées", "Non"],
            ["min_area_m2", "Surface minimale des détections conservées (m²)", "0"],
        ],
        col_widths=[4, 7, 3.5]
    )

    # B.3
    add_heading2(doc, "B.3 Post-traitement des détections")
    add_body(doc, "Après l'inférence, les détections subissent automatiquement :")
    add_bullet(doc, "Validation géométrique — correction des géométries invalides")
    add_bullet(doc, "Fusion intra-classe — les détections de la même classe qui se chevauchent sont fusionnées")
    add_bullet(doc, "Suppression des recouvrements inter-classes — en cas de recouvrement, la détection avec la confiance la plus faible est supprimée")
    add_bullet(doc, "Filtre de surface minimale — les détections plus petites que min_area_m2 sont supprimées")

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════
    # ANNEXE C — CLUSTERING SPATIAL (DBSCAN)
    # ═══════════════════════════════════════════════════════

    add_heading1(doc, "Annexe C — Clustering spatial (DBSCAN)")

    add_body(doc,
        "Le clustering permet de regrouper des détections individuelles en zones "
        "cohérentes. Par exemple, regrouper des cratères d'obus individuels en "
        "une zone de bombardement."
    )
    add_body(doc,
        "Cette fonctionnalité est configurée par modèle dans le fichier args.yaml "
        "de chaque modèle."
    )

    add_heading2(doc, "C.1 Principe")
    add_body(doc,
        "L'algorithme DBSCAN (Density-Based Spatial Clustering) identifie des groupes "
        "de détections en fonction de leur proximité spatiale et de leur densité. "
        "Les détections isolées (hors de tout groupe) sont ignorées pour le clustering "
        "mais conservées individuellement."
    )

    add_heading2(doc, "C.2 Paramètres de configuration (args.yaml)")
    add_code_block(doc,
        "clustering:\n"
        "  - target_classes: [\"cratere_obus\"]   # Classes sources à regrouper\n"
        "    min_confidence: 0.4                # Confiance minimale pour participer\n"
        "    eps_m: 60                          # Distance max entre détections (mètres)\n"
        "    min_cluster_size: 10               # Nb minimum de détections par groupe\n"
        "    min_samples: 5                     # Paramètre DBSCAN : voisinage minimum\n"
        "    output_class_name: \"zone_crateres\" # Nom de la classe de sortie\n"
        "    output_geometry: \"convex_hull\"     # convex_hull ou bounding_box\n"
        "    buffer_m: 10                       # Marge autour de la géométrie (mètres)\n"
        "    min_area_m2: 500                   # Surface minimale conservée (m²)"
    )

    add_heading2(doc, "C.3 Description des paramètres")
    add_styled_table(doc,
        ["Paramètre", "Description"],
        [
            ["target_classes", "Liste des classes de détections individuelles à regrouper"],
            ["min_confidence", "Les détections sous ce seuil sont exclues du clustering"],
            ["eps_m", "Distance en mètres en dessous de laquelle deux détections sont voisines. Paramètre clé à calibrer"],
            ["min_cluster_size", "Un groupe avec moins de détections que cette valeur est ignoré"],
            ["min_samples", "Nombre minimum de voisins dans le rayon eps_m pour qu'un point soit considéré dense"],
            ["output_class_name", "Nom de la classe de sortie dans les shapefiles"],
            ["output_geometry", "convex_hull = enveloppe convexe, bounding_box = rectangle englobant"],
            ["buffer_m", "Marge ajoutée autour de la géométrie finale"],
            ["min_area_m2", "Les zones inférieures à cette surface sont supprimées"],
        ],
        col_widths=[4.5, 10]
    )

    add_heading2(doc, "C.4 Pondération par confiance (confidence_weight)")
    add_body(doc,
        "Un paramètre optionnel confidence_weight (défaut : 0) module la distance "
        "effective entre détections en fonction de leur confiance respective. "
        "Deux détections très confiantes seront considérées plus proches."
    )
    add_code_block(doc,
        "confidence_weight: 0.5   # 0 = DBSCAN classique, >0 = pondération par confiance"
    )

    add_heading2(doc, "C.5 Attributs des shapefiles de zones")
    add_styled_table(doc,
        ["Attribut", "Description"],
        [
            ["class_name", "Nom de la classe de zone"],
            ["confidence", "Confiance moyenne des détections du groupe"],
            ["nb_detect", "Nombre de détections individuelles dans la zone"],
            ["area_m2", "Surface de la zone en m²"],
            ["density", "Densité de détections par hectare"],
            ["cluster_id", "Identifiant unique du groupe"],
        ],
        col_widths=[4.5, 10]
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════
    # GLOSSAIRE
    # ═══════════════════════════════════════════════════════

    add_heading1(doc, "📚  Glossaire")

    add_body(doc, "Lexique des termes techniques utilisés dans cette documentation.")

    add_styled_table(doc,
        ["Terme", "Définition"],
        [
            ["ASPRS", "American Society for Photogrammetry and Remote Sensing — classification standard des points LiDAR (2 = sol, 6 = bâtiment…)."],
            ["COPC", "Cloud-Optimized Point Cloud — format LAZ indexé par l'IGN pour un accès rapide."],
            ["CV", "Computer Vision — détection automatique d'entités par intelligence artificielle."],
            ["DBSCAN", "Density-Based Spatial Clustering — algorithme de regroupement spatial fondé sur la densité."],
            ["GeoTIFF", "Format raster TIF avec géoréférencement embarqué."],
            ["GPKG", "GeoPackage — format de données vectorielles (équivalent moderne du shapefile)."],
            ["IGN LiDAR HD", "Programme national IGN fournissant des données LiDAR haute densité sur la France métropolitaine."],
            ["IoU", "Intersection over Union — mesure de recouvrement entre deux polygones (0 à 1)."],
            ["LAZ / LAS", "Formats de nuages de points LiDAR (LAZ est une version compressée de LAS)."],
            ["LD", "Local Dominance — indice RVT mesurant la dominance locale d'un point."],
            ["LiDAR", "Light Detection And Ranging — technologie de télémétrie laser."],
            ["M-HS", "Multi-directional Hillshade — ombrage multidirectionnel."],
            ["MNT", "Modèle Numérique de Terrain — raster d'altitude du sol nu."],
            ["MNS", "Modèle Numérique de Surface — raster d'altitude incluant le sursol (arbres, bâtiments…)."],
            ["ONNX", "Open Neural Network Exchange — format standard de modèles d'IA."],
            ["PDAL", "Point Data Abstraction Library — bibliothèque de traitement des nuages de points LiDAR."],
            ["RVT", "Relief Visualization Toolbox — ensemble d'indices de visualisation du relief."],
            ["SAHI", "Slicing Aided Hyper Inference — technique de découpage d'images en tuiles pour la détection."],
            ["Shapefile", "Format vectoriel .shp (ESRI) — souvent remplacé ici par GeoPackage."],
            ["SLO", "Slope — carte de pente du terrain."],
            ["SLRM", "Simple Local Relief Model — isole le micro-relief en soustrayant une surface lissée au MNT."],
            ["SVF", "Sky View Factor — proportion de ciel visible depuis chaque point (révèle les creux)."],
            ["VAT", "Visualisation Archéologique Totale — combinaison d'indices RVT optimisée pour l'archéologie."],
            ["VRT", "Virtual Raster — fichier XML qui assemble virtuellement plusieurs rasters en mosaïque."],
        ],
        col_widths=[3, 12]
    )

    return doc


# ─── POINT D'ENTRÉE ─────────────────────────────────────────

if __name__ == "__main__":
    import os
    from pathlib import Path

    output_dir = Path(__file__).parent
    output_path = output_dir / "documentation_utilisateur.docx"

    print("Génération de la documentation...")
    doc = build_document()
    doc.save(str(output_path))
    print(f"Documentation générée : {output_path}")
    os.startfile(str(output_path))
    